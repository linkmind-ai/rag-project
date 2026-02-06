import asyncio
from pathlib import Path

import aiofiles
from common.config import settings
from models.state import Document


class FileProcessor:
    """파일 처리 클래스"""

    def __init__(self) -> None:
        self._upload_dir = Path(settings.UPLOAD_DIR)
        self._upload_dir.mkdir(exist_ok=True)

    async def save_file(self, filename: str, content: bytes) -> str:
        """파일 저장"""
        file_path = self._upload_dir / filename

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        return str(file_path)

    def validate_files(self, filename: str, file_size: int) -> tuple[bool, str | None]:
        if file_size > settings.MAX_FILE_SIZE:
            return False, "파일 크기가 너무 큽니다."

        ext = filename.split(".")[-1].lower()
        if ext not in settings.ALLOWED_EXTENSIONS:
            return False, "지원하지 않는 파일 형식입니다."

        return True, None

    async def process_txt_file(self, file_path: str) -> list[Document]:
        """텍스트 파일 처리 프로세스"""
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        filename = Path(file_path).name
        return [
            Document(
                content=content,
                metadata={
                    "source": filename,
                    "file_type": "txt",
                    "file_path": file_path,
                },
            )
        ]

    async def process_pdf_file(self, file_path: str) -> list[Document]:
        """PDF 파일 처리 프로세스"""
        try:
            import pypdf

            documents = []

            async def read_pdf() -> None:
                with open(file_path, "rb") as f:
                    pdf_reader = pypdf.PdfReader(f)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            documents.append(
                                Document(
                                    content=text,
                                    metadata={
                                        "source": Path(file_path).name,
                                        "file_type": "pdf",
                                        "page_number": page_num + 1,
                                        "file_path": file_path,
                                    },
                                )
                            )

            await asyncio.to_thread(read_pdf)
            return documents

        except ImportError as e:
            raise ImportError("pypdf 실행 중 오류 발생") from e

    async def process_docx_file(self, file_path: str) -> list[Document]:
        """DOCX 파일 처리 프로세스"""
        try:
            import docx

            async def read_docx() -> str:
                doc = docx.Document(file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                return "\n\n".join(paragraphs)

            content = await asyncio.to_thread(read_docx)

            filename = Path(file_path).name
            return [
                Document(
                    content=content,
                    metadata={
                        "source": filename,
                        "file_type": "txt",
                        "file_path": file_path,
                    },
                )
            ]

        except ImportError as e:
            raise ImportError("python-docx 실행 오류 발생") from e

    async def process_md_file(self, file_path: str) -> list[Document]:
        """마크다운 처리 프로세스"""
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        filename = Path(file_path).name
        return [
            Document(
                content=content,
                metadata={
                    "source": filename,
                    "file_type": "markdown",
                    "file_path": file_path,
                },
            )
        ]

    async def process_file(self, file_path: str) -> list[Document]:
        """파일 확장자별 프로세서 결정"""
        ext = Path(file_path).suffix.lower()[1:]

        processors = {
            "txt": self.process_txt_file,
            "pdf": self.process_pdf_file,
            "docx": self.process_docx_file,
            "md": self.process_md_file,
        }

        processor = processors.get(ext)
        if not processor:
            raise ValueError("지원하지 않는 파일 형식: {ext}")

        return await processor(file_path)

    async def delete_file(self, file_path: str) -> bool:
        """파일 삭제"""
        try:
            path = Path(file_path)
            if path.exists():
                await asyncio.to_thread(path.unlink)
                return True
            return False
        except Exception:
            return False


file_processor = FileProcessor()
